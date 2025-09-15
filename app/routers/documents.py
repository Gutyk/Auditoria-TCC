from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from ..db import SessionLocal
from ..models import Document, Project
from ..schemas import (
    DocumentIn,
    DocumentOut,
    DocumentDetailOut,
    DocumentUpdateIn,
)

# ---- opcional: parsers p/ upload ----
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    import docx  # python-docx
except Exception:
    docx = None

router = APIRouter()

# =============================
# DB dependency
# =============================
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# =============================
# Helpers (para upload opcional)
# =============================
def _extract_text_from_bytes(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".txt"):
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            return ""
    if name.endswith(".pdf"):
        if not PdfReader:
            raise HTTPException(status_code=500, detail="Dependência ausente: pypdf")
        import io
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for pg in reader.pages:
            try:
                pages.append(pg.extract_text() or "")
            except Exception:
                pages.append("")
        return "\n".join(pages)
    if name.endswith(".docx"):
        if not docx:
            raise HTTPException(status_code=500, detail="Dependência ausente: python-docx")
        import io
        d = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in d.paragraphs)

    # fallback: tenta como texto puro
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return ""

# =============================
# Rotas
# =============================

@router.post("", response_model=DocumentOut)
def create_document(payload: DocumentIn, db: Session = Depends(get_db)):
    if not db.query(Project).filter_by(id=payload.project_id).first():
        raise HTTPException(status_code=404, detail="Project not found")
    d = Document(project_id=payload.project_id, title=payload.title, content=payload.content)
    db.add(d)
    db.commit()
    db.refresh(d)
    return d

@router.get("/{project_id}", response_model=List[DocumentOut])
def list_documents(project_id: int, db: Session = Depends(get_db)):
    return (
        db.query(Document)
        .filter_by(project_id=project_id)
        .order_by(Document.id.desc())
        .all()
    )

@router.get("/detail/{doc_id}", response_model=DocumentDetailOut)
def get_document(doc_id: int, db: Session = Depends(get_db)):
    d = db.query(Document).filter_by(id=doc_id).first()
    if not d:
        raise HTTPException(status_code=404, detail="Document not found")
    return d

@router.put("/{doc_id}", response_model=DocumentDetailOut)
def update_document(doc_id: int, payload: DocumentUpdateIn, db: Session = Depends(get_db)):
    d = db.query(Document).filter_by(id=doc_id).first()
    if not d:
        raise HTTPException(status_code=404, detail="Document not found")
    if payload.title is not None:
        d.title = payload.title
    if payload.content is not None:
        d.content = payload.content
    db.add(d)
    db.commit()
    db.refresh(d)
    return d

# (Opcional) Upload direto na API — só use se o front for enviar multipart para /documents/upload
@router.post("/upload", response_model=DocumentOut)
def upload_document(
    project_id: int = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    if not db.query(Project).filter_by(id=project_id).first():
        raise HTTPException(status_code=404, detail="Project not found")

    raw = file.file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Arquivo vazio")

    content = _extract_text_from_bytes(file.filename, raw).strip()
    if not content:
        raise HTTPException(status_code=400, detail="Não foi possível extrair texto do arquivo")

    d = Document(project_id=project_id, title=file.filename, content=content)
    db.add(d)
    db.commit()
    db.refresh(d)
    return d


@router.delete("/{document_id}")
def delete_document(document_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento não encontrado")
    db.delete(doc)
    db.commit()
    return {"message": f"Documento {document_id} deletado com sucesso"}
